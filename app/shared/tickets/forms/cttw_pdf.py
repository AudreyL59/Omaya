"""Régénération du PDF de contrat W signé (FI_CttW Plan 2).

Reproduit la chaîne WinDev (code init Plan 2 + AjouterSignatures) :
  1. extrait le mémo binaire TK_DemandeCttW.Contenu (docx)
  2. insère paraphe (pied de page), signature (balises S_SIGN),
     mention « Lu & approuvé » (balises S_MENTION) — images prises
     dans les mémos, sinon récupérées en HTTP (rest.omaya.fr/sign).
  3. docx -> PDF via LibreOffice headless
  4. génère l'état récap (EtatRecapSignCttW) en reportlab
  5. fusionne contrat + récap

Le PDF final est servi (Plan 2) et uploadé dans le dossier salarié
(« Ce contrat de travail est valide »).
"""

from __future__ import annotations

import base64
import io
import os
import subprocess
import tempfile
import uuid

import urllib.request

from app.core.config import (
    CTTW_SIGN_URL,
    CTTW_SIGN_URL_FALLBACK,
    SOFFICE_BIN,
)
from app.core.database import get_connection

CREATE_NO_WINDOW = 0x08000000


# ---------------------------------------------------------------
# Lecture mémos binaires (le bridge encode les binaires en base64 ;
# on SELECT toujours la clé + le mémo — un SELECT du seul mémo
# binaire casse le JSON du bridge).
# ---------------------------------------------------------------

def _memo_bytes(id_ticket: int, field: str) -> bytes | None:
    try:
        db = get_connection("ticket_rh")
        r = db.query_one(
            f"SELECT IDTK_Liste, {field} FROM TK_DemandeCttW "
            f"WHERE IDTK_Liste = ?",
            (int(id_ticket),),
        )
        v = r.get(field) if r else None
        if not v:
            return None
        if isinstance(v, bytes):
            return v
        return base64.b64decode(v)
    except Exception:
        return None


def ftp_upload(remote_dir: str, filename: str, data: bytes) -> None:
    """Upload `data` vers <remote_dir>/<filename> (crée l'arbo au besoin).
    remote_dir absolu, ex: /OMAYA/gestionRH/<id>/Fiches_Salaires."""
    import ftplib

    from app.core.config import FTP_HOST, FTP_PASSWORD, FTP_USER

    ftp = ftplib.FTP(timeout=30)
    ftp.encoding = "latin-1"
    ftp.connect(FTP_HOST, 21)
    ftp.login(FTP_USER, FTP_PASSWORD)
    try:
        ftp.cwd("/")
        for part in [p for p in remote_dir.split("/") if p]:
            try:
                ftp.cwd(part)
            except ftplib.error_perm:
                ftp.mkd(part)
                ftp.cwd(part)
        ftp.storbinary(f"STOR {filename}", io.BytesIO(data))
    finally:
        try:
            ftp.quit()
        except Exception:
            pass


def _http_image(url: str) -> bytes | None:
    try:
        with urllib.request.urlopen(url, timeout=15) as resp:
            data = resp.read()
        return data if data and len(data) > 64 else None
    except Exception:
        return None


def _image_or_fallback(id_ticket: int, field: str, suffix: str) -> bytes | None:
    """Mémo binaire `field` sinon HTTP rest.omaya.fr puis sos.rest."""
    img = _memo_bytes(id_ticket, field)
    if img:
        return img
    for base in (CTTW_SIGN_URL, CTTW_SIGN_URL_FALLBACK):
        img = _http_image(f"{base}/{id_ticket}_{suffix}.jpg")
        if img:
            return img
    return None


# ---------------------------------------------------------------
# Manipulation docx (python-docx)
# ---------------------------------------------------------------

def _iter_paragraphs(doc):
    """Tous les paragraphes : corps + tables + en-têtes/pieds."""
    from docx.document import Document as _Doc  # noqa

    def _para_in(parent):
        for p in parent.paragraphs:
            yield p
        for t in parent.tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from _para_in(cell)

    yield from _para_in(doc)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer):
            try:
                yield from _para_in(hf)
            except Exception:
                continue


def _replace_token_with_image(doc, token: str, img: bytes, height_mm: float):
    """Remplace le `token` par l'image, EN PLACE dans le run qui le
    contient, sans toucher au reste du paragraphe (préserve l'espace
    réservé et donc la pagination — sinon les ~125 images ancrées
    logo/paraphe se désalignent et s'empilent).

    L'image est dimensionnée par sa HAUTEUR (la largeur suit le ratio)
    pour rester dans la zone de signature réservée.
    """
    from docx.shared import Mm

    for p in _iter_paragraphs(doc):
        if token not in p.text:
            continue
        for run in list(p.runs):
            if token not in run.text:
                continue
            before, _, after = run.text.partition(token)
            run.text = before
            try:
                run.add_picture(io.BytesIO(img), height=Mm(height_mm))
            except Exception:
                pass
            if after:
                # remet le texte qui suivait le token
                tail = p.add_run(after)
                tail.font.size = run.font.size


def _add_paraphe_footer(doc, img: bytes):
    """Ajoute le paraphe dans le pied de page (aligné à droite).

    /!\\ Les sections « liées au précédent » partagent le MÊME
    footerN.xml : itérer toutes les sections et appeler add_paragraph()
    insère le paraphe autant de fois qu'il y a de sections dans ce
    pied unique -> logo en double/triple sur chaque page.
    On insère donc une seule fois par pied de page DISTINCT (on saute
    les pieds liés, qui héritent automatiquement du paraphe ajouté au
    pied propriétaire)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    seen: set[int] = set()
    for section in doc.sections:
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                if footer is None or footer.is_linked_to_previous:
                    continue
                key = id(footer._element)
            except Exception:
                continue
            if key in seen:
                continue
            seen.add(key)
            para = footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                para.add_run().add_picture(io.BytesIO(img), width=Mm(12))
            except Exception:
                pass


def _strip_floating_anchors(doc):
    """Supprime du CORPS toutes les images flottantes (`<wp:anchor>`).

    Le moteur d'état WinDev tamponne le logo + le paraphe en BAS de
    CHAQUE page sous forme d'images ancrées dans le corps (toutes à la
    même position : posOffset V identique, relativeFrom=margin). WinDev
    paginant différemment de LibreOffice, ces ancres se regroupent sur
    les pages denses -> logo en double/triple en bas à gauche.

    Le vrai pied de page Word (footer1.xml) porte déjà le logo + le
    paraphe + le n° de page : il se répète proprement une fois par
    page. On supprime donc les ancres redondantes du corps et on laisse
    le pied de page faire le travail. Les images `<wp:inline>` (tampons
    de signature dans le contenu) et les balises S_SIGN/S_MENTION ne
    sont pas touchées.
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    to_remove = []
    for drawing in body.iter(qn("w:drawing")):
        if drawing.find(qn("wp:anchor")) is not None:
            to_remove.append(drawing)
    for drawing in to_remove:
        parent = drawing.getparent()  # le <w:r>
        if parent is not None:
            gp = parent.getparent()
            if gp is not None:
                gp.remove(parent)
            else:
                parent.remove(drawing)
    return len(to_remove)


# Ancre paraphe salarié : clone de l'ancre du paraphe gérant
# (footer1.xml rId1 : posH 5760000, posV 8640000, 360000x360000),
# décalée à gauche pour être « à côté » de celle du gérant.
_PARAPHE_SAL_RUN = (
    '<w:r><w:drawing>'
    '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/'
    '2006/wordprocessingDrawing" relativeHeight="3" simplePos="0" '
    'behindDoc="0" locked="0" layoutInCell="0" hidden="0" '
    'allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="margin"><wp:posOffset>5220000'
    '</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="margin"><wp:posOffset>8640000'
    '</wp:posOffset></wp:positionV>'
    '<wp:extent cx="360000" cy="360000"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
    '<wp:docPr id="990" name="ParapheSalarie"/>'
    '<wp:cNvGraphicFramePr/>'
    '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/'
    '2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/'
    'drawingml/2006/picture">'
    '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/'
    '2006/picture"><pic:nvPicPr>'
    '<pic:cNvPr id="990" name="ParapheSalarie"/><pic:cNvPicPr/>'
    '</pic:nvPicPr><pic:blipFill>'
    '<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/'
    '2006/relationships" r:embed="__RID__" cstate="none"/><a:stretch/>'
    '</pic:blipFill><pic:spPr><a:xfrm><a:ext cx="360000" cy="360000"/>'
    '</a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic>'
    '</a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>'
)


def _inject_salarie_paraphe_footer(docx_path: str, paraphe: bytes) -> None:
    """Ajoute le paraphe du salarié dans footer1.xml (3ᵉ image ancrée,
    à gauche de celle du gérant). Chirurgie directe du zip docx :
    python-docx ne gère pas les images ancrées (wp:anchor)."""
    import re
    import shutil
    import zipfile

    ext = "png" if paraphe[:4] == b"\x89PNG" else "jpg"
    media_name = f"word/media/parapheSal.{ext}"
    rels_name = "word/_rels/footer1.xml.rels"
    footer_name = "word/footer1.xml"

    zin = zipfile.ZipFile(docx_path, "r")
    names = zin.namelist()
    if footer_name not in names or rels_name not in names:
        zin.close()
        return

    rels = zin.read(rels_name).decode("utf-8", "ignore")
    nums = [int(n) for n in re.findall(r'Id="rId(\d+)"', rels)]
    new_rid = f"rId{(max(nums) + 1) if nums else 1}"
    new_rel = (
        f'<Relationship Id="{new_rid}" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
        'relationships/image" '
        f'Target="media/parapheSal.{ext}"/>'
    )
    rels = rels.replace("</Relationships>", new_rel + "</Relationships>")

    footer = zin.read(footer_name).decode("utf-8", "ignore")
    run = _PARAPHE_SAL_RUN.replace("__RID__", new_rid)
    # insère le run juste avant la fin du 1er paragraphe (celui qui
    # contient déjà les ancres logo + paraphe gérant)
    footer = footer.replace("</w:p>", run + "</w:p>", 1)

    tmp_path = docx_path + ".tmp"
    zout = zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED)
    for item in zin.infolist():
        if item.filename in (footer_name, rels_name):
            continue
        zout.writestr(item, zin.read(item.filename))
    zout.writestr(footer_name, footer)
    zout.writestr(rels_name, rels)
    zout.writestr(media_name, paraphe)
    zout.close()
    zin.close()
    shutil.move(tmp_path, docx_path)


# ---------------------------------------------------------------
# Conversion docx -> PDF (LibreOffice headless)
# ---------------------------------------------------------------

class SofficeError(Exception):
    pass


def _docx_to_pdf(docx_path: str, out_dir: str) -> str:
    if not os.path.exists(SOFFICE_BIN):
        raise SofficeError(
            f"LibreOffice introuvable ({SOFFICE_BIN}). Installer LibreOffice "
            "ou définir SOFFICE_BIN."
        )
    proc = subprocess.run(
        [SOFFICE_BIN, "--headless", "--norestore", "--convert-to", "pdf",
         "--outdir", out_dir, docx_path],
        stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        timeout=120, creationflags=CREATE_NO_WINDOW,
    )
    pdf_path = os.path.join(
        out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    )
    if not os.path.exists(pdf_path):
        raise SofficeError(
            f"Conversion PDF échouée (exit {proc.returncode}): "
            f"{proc.stderr.decode('utf-8', 'ignore')[:300]}"
        )
    return pdf_path


# ---------------------------------------------------------------
# ContenuValidation : "<page>\t<date>\n...//<codeSMS>"
# ---------------------------------------------------------------

def _fmt_dt(v: str) -> str:
    """WinDev compact AAAAMMJJHHMMSS[mmm] -> 'JJ/MM/AAAA HH:MM:SS'.
    Laisse tel quel si déjà formaté."""
    d = "".join(c for c in str(v or "") if c.isdigit())
    if len(d) >= 14:
        return f"{d[6:8]}/{d[4:6]}/{d[0:4]} {d[8:10]}:{d[10:12]}:{d[12:14]}"
    return str(v or "")


def parse_contenu_validation(s: str) -> tuple[list[tuple[str, str]], str]:
    s = s or ""
    code_sms = ""
    pages_part = s
    if "//" in s:
        pages_part, code_sms = s.split("//", 1)
    pages: list[tuple[str, str]] = []
    for line in pages_part.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = line.strip()
        if not line:
            continue
        if "\t" in line:
            a, b = line.split("\t", 1)
        else:
            parts = line.split(None, 1)
            a, b = (parts + ["", ""])[:2]
        pages.append((a.strip(), b.strip()))
    return pages, code_sms.strip()


# ---------------------------------------------------------------
# État récap EtatRecapSignCttW (page finale)
# ---------------------------------------------------------------

def _build_recap_pdf(
    photo: bytes | None,
    signature: bytes | None,
    lu_app: bytes | None,
    date_sign: str,
    code_sms: str,
    salarie_nom: str,
    da_nom: str,
    pages: list[tuple[str, str]],
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    c.setFont("Helvetica-Bold", 17)
    c.drawCentredString(w / 2, h - 25 * mm,
                        "Signature dématérialisée du CONTRAT de TRAVAIL")

    y = h - 45 * mm
    c.setFont("Helvetica", 10)
    c.drawString(25 * mm, y, f"Ce contrat de travail a été signé le :  {date_sign}")

    # Photo signataire (gauche)
    if photo:
        try:
            c.drawImage(ImageReader(io.BytesIO(photo)), 25 * mm, y - 38 * mm,
                        28 * mm, 34 * mm, preserveAspectRatio=True, mask="auto")
        except Exception:
            pass
    tx = 60 * mm
    c.drawString(tx, y - 8 * mm, f", en présence de :  {da_nom}")
    c.drawString(tx, y - 15 * mm, f", par :  {salarie_nom}")
    c.setFont("Helvetica", 7.5)
    c.drawString(tx, y - 23 * mm,
                 "Photographie du signataire prise au moment de la signature "
                 "électronique.")
    c.drawString(tx, y - 27 * mm,
                 "Cette photographie est confidentielle, elle ne sera pas "
                 "diffusée sans autorisation écrite du salarié.")

    # Tableau pages validées
    ty = y - 50 * mm
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(75 * mm, ty, "Page")
    c.drawCentredString(135 * mm, ty, "Date et heure Validation")
    c.setFont("Helvetica", 9)
    ty -= 6 * mm
    for num, dte in pages:
        if ty < 45 * mm:
            c.showPage()
            ty = h - 25 * mm
            c.setFont("Helvetica", 9)
        c.drawCentredString(75 * mm, ty, str(num))
        c.drawCentredString(135 * mm, ty, _fmt_dt(dte))
        ty -= 5 * mm

    ty -= 4 * mm
    c.setFont("Helvetica", 9)
    if code_sms:
        # code_sms contient déjà la phrase complète
        # (« Vérification SMS envoyée sur le …, code … »)
        c.drawString(25 * mm, ty, code_sms[:120])
        ty -= 8 * mm
    c.drawString(25 * mm, ty, "Mention « Lu et approuvé » :")
    c.drawString(115 * mm, ty, "Signature :")
    if lu_app:
        try:
            c.drawImage(ImageReader(io.BytesIO(lu_app)), 25 * mm, ty - 28 * mm,
                        60 * mm, 25 * mm, preserveAspectRatio=True, mask="auto")
        except Exception:
            pass
    if signature:
        try:
            c.drawImage(ImageReader(io.BytesIO(signature)), 115 * mm,
                        ty - 28 * mm, 55 * mm, 25 * mm,
                        preserveAspectRatio=True, mask="auto")
        except Exception:
            pass

    c.setFont("Helvetica-Bold", 8)
    c.drawString(20 * mm, 15 * mm, "Récapitulatif Signature Dématérialisée")
    c.drawRightString(w - 20 * mm, 15 * mm, "1/1")
    c.showPage()
    c.save()
    return buf.getvalue()


# ---------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------

def regenerate_signed_pdf(
    id_ticket: int,
    salarie_nom: str,
    da_nom: str,
    date_sign: str,
) -> bytes:
    """Régénère le PDF signé complet (contrat + signatures + récap)."""
    from docx import Document
    from pypdf import PdfReader, PdfWriter

    docx_raw = _memo_bytes(id_ticket, "Contenu")
    if not docx_raw:
        raise ValueError("Contrat (Contenu) introuvable ou vide")

    sign = _image_or_fallback(id_ticket, "Signature", "CttWSignature")
    lu_app = _image_or_fallback(id_ticket, "luApp", "CttWLuApp")
    paraphe = _memo_bytes(id_ticket, "paraphe")
    photo = _memo_bytes(id_ticket, "PhotoSalarié")

    # ContenuValidation (pages validées + code SMS)
    db = get_connection("ticket_rh")
    cv = db.query_one(
        "SELECT IDTK_Liste, ContenuValidation FROM TK_DemandeCttW "
        "WHERE IDTK_Liste = ?",
        (int(id_ticket),),
    )
    pages, code_sms = parse_contenu_validation(
        (cv.get("ContenuValidation") if cv else "") or ""
    )

    tmp = tempfile.mkdtemp(prefix=f"cttw_{id_ticket}_")
    docx_path = os.path.join(tmp, f"{id_ticket}.docx")
    with open(docx_path, "wb") as f:
        f.write(docx_raw)

    doc = Document(docx_path)
    # Supprime les logos/paraphes ancrés par page dans le corps
    # (dupliqués par la repagination LibreOffice). Le pied de page
    # Word (footer1.xml) porte déjà logo + paraphe + n° de page et se
    # répète proprement une fois par page -> on ne réajoute rien.
    _strip_floating_anchors(doc)
    if sign:
        _replace_token_with_image(doc, "S_SIGN", sign, 18)
    if lu_app:
        _replace_token_with_image(doc, "S_MENTION", lu_app, 16)
    doc.save(docx_path)

    # Paraphe du salarié dans le pied de page, à côté du gérant
    # (footer1.xml -> rendu 1x/page nativement, pas de duplication).
    if paraphe:
        try:
            _inject_salarie_paraphe_footer(docx_path, paraphe)
        except Exception:
            pass

    pdf_contrat = _docx_to_pdf(docx_path, tmp)
    recap_pdf = _build_recap_pdf(
        photo, sign, lu_app, date_sign, code_sms,
        salarie_nom, da_nom, pages,
    )

    writer = PdfWriter()
    for src in (open(pdf_contrat, "rb").read(), recap_pdf):
        reader = PdfReader(io.BytesIO(src))
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()
