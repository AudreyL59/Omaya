"""Publipostage DOCX pour les contrats de courtage.

Reproduit les procedures WinDev :
  - Publipostage_Distrib : tags STE_..._DISTRIB (societe distrib +
    infos gerant salarie S_...)
  - Publipostage_Rem : insertion d'un tableau REM apres chaque tag
    TABLEAU_REM (une iteration par GroupeRem actif du distrib+groupe_op)
  - Publipostage_STE : tags STE_... (societe editrice = id_ste du doc)

Version 1 : remplacement de tous les tags texte + tableaux REM.
Non implemente pour l'instant :
  - Images (STE_LOGO, GER_SIGN, STE_CACHET, GerantParaphe) : necessite
    manipulation avancee de python-docx + gestion du buffer bytea
  - Pied de page : idem
  - Conversion PDF (necessite docx2pdf ou LibreOffice)
Ces fonctionnalites viendront dans un commit ulterieur si besoin.
"""

from __future__ import annotations

import io
from datetime import date, datetime
from typing import Any

from app.core.database.pg import get_pg_connection

GROUPE_OP_AUTRE = 281474976710657   # "Autres" : pas de publipostage REM/STE


def _fmt_date_fr(v: Any) -> str:
    if v is None: return ""
    if isinstance(v, (date, datetime)):
        return v.strftime("%d/%m/%Y")
    s = str(v)
    if len(s) >= 10 and s[4] == "-":
        return f"{s[8:10]}/{s[5:7]}/{s[0:4]}"
    return s


def _get_distrib_info(id_distrib: int) -> dict[str, str]:
    """Recupere infos societe distrib + gerant salarie. Cf reqinfoDistri WinDev."""
    db = get_pg_connection("rh")
    r = db.query_one(
        """SELECT s.raison_sociale, s.code_ape, s.rcs, s.capital,
                  s.adresse1, s.cp, s.ville,
                  s.gerant_nom, s.gerant_type, s.siret, s.siren, s.id_gerant,
                  fj.lib_form_juri,
                  sa.civilite, sa.nom AS sa_nom, sa.prenom AS sa_prenom,
                  sa.lieu_naiss, sa.dep_naiss, sa.num_ss, sa.date_naiss
             FROM rh.pgt_societe s
             LEFT JOIN rh.pgt_societe_formjuri fj
                    ON CAST(fj.id_societe_form_juri AS text) = s.forme_juri
             LEFT JOIN rh.pgt_salarie sa ON sa.id_salarie = s.id_gerant
            WHERE s.id_ste = ? LIMIT 1""",
        (int(id_distrib),),
    ) or {}
    civ = int(r.get("civilite") or 0)
    s_titre = "Mme" if civ == 2 else "Mr."
    return {
        "STE_RS_DISTRIB": r.get("raison_sociale") or "",
        "STE_APE_DISTRIB": r.get("code_ape") or "",
        "STE_RCS_DISTRIB": r.get("rcs") or "",
        "STE_FORMJURI_DISTRIB": r.get("lib_form_juri") or "",
        "STE_CAPITAL_DISTRIB": f"{r.get('capital') or 0}",
        "STE_ADR_DISTRIB": " ".join(filter(None, [
            r.get("adresse1"), r.get("cp"), r.get("ville"),
        ])),
        "STE_ADRESSE_DISTRIB": r.get("adresse1") or "",
        "STE_CP_DISTRIB": r.get("cp") or "",
        "STE_VILLE_DISTRIB": r.get("ville") or "",
        "STE_GERANT_NOM_DISTRIB": r.get("gerant_nom") or "",
        "STE_GERANT_TYPE_DISTRIB": r.get("gerant_type") or "",
        "STE_SIRET_DISTRIB": r.get("siret") or "",
        "STE_SIREN_DISTRIB": r.get("siren") or "",
        # Tags salarie gerant (S_...)
        "S_TITRE": s_titre,
        "S_NOM": r.get("sa_nom") or "",
        "S_PRENOM": r.get("sa_prenom") or "",
        "S_LNAISS": r.get("lieu_naiss") or "",
        "S_DEPNAISS": str(r.get("dep_naiss") or ""),
        "S_NUMSS": r.get("num_ss") or "",
        "S_DNAISS": _fmt_date_fr(r.get("date_naiss")),
    }


def _get_ste_info(id_ste: int) -> dict[str, str]:
    """Cf reqInfoSTe Publipostage_STE : infos societe editrice."""
    if not id_ste: return {}
    db = get_pg_connection("rh")
    r = db.query_one(
        """SELECT raison_sociale, code_ape, rcs, capital,
                  adresse1, cp, ville, gerant_nom, gerant_type,
                  siret, siren
             FROM rh.pgt_societe WHERE id_ste = ? LIMIT 1""",
        (int(id_ste),),
    ) or {}
    return {
        "STE_RS": r.get("raison_sociale") or "",
        "STE_APE": r.get("code_ape") or "",
        "STE_RCS": r.get("rcs") or "",
        "STE_CAPITAL": f"{r.get('capital') or 0}",
        "STE_ADR": " ".join(filter(None, [
            r.get("adresse1"), r.get("cp"), r.get("ville"),
        ])),
        "STE_VILLE": r.get("ville") or "",
        "STE_GERANT_NOM": r.get("gerant_nom") or "",
        "STE_GERANT_TYPE": r.get("gerant_type") or "",
        "STE_SIRET": r.get("siret") or "",
        "STE_SIREN": r.get("siren") or "",
    }


def _get_groupes_rem_publipostage(
    id_distrib: int, id_groupe_operateur: int,
) -> list[dict]:
    """Cf reqListeRem Publipostage_Rem : liste des GroupeRem actifs
    (par distrib + groupe operateur) avec leurs grilles X/Y/Tab
    pre-fabriquees pour affichage tableau."""
    db = get_pg_connection("adv")
    grs = db.query(
        """SELECT id_groupe_rem, lib_groupe, nb_col, nb_ligne, ordre
             FROM adv.pgt_groupe_rem
            WHERE id_groupe_operateur = ?
              AND id_distrib = ?
              AND is_actif = TRUE
              AND (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
            ORDER BY ordre DESC""",
        (int(id_groupe_operateur), int(id_distrib)),
    ) or []
    out = []
    for gr in grs:
        id_gr = int(gr["id_groupe_rem"])
        xs = db.query(
            """SELECT id_groupe_rem_x, lib, code_interne
                 FROM adv.pgt_groupe_rem_x
                WHERE id_groupe_rem = ?
                  AND (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
                ORDER BY ordre""",
            (id_gr,),
        ) or []
        ys = db.query(
            """SELECT id_groupe_rem_y, lib, code_interne
                 FROM adv.pgt_groupe_rem_y
                WHERE id_groupe_rem = ?
                  AND (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
                ORDER BY ordre""",
            (id_gr,),
        ) or []
        cells = db.query(
            """SELECT id_groupe_rem_x, id_groupe_rem_y, montant
                 FROM adv.pgt_groupe_rem_tab
                WHERE id_groupe_rem = ?
                  AND (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')""",
            (id_gr,),
        ) or []
        cell_map = {(int(c["id_groupe_rem_x"]), int(c["id_groupe_rem_y"])): float(c["montant"] or 0) for c in cells}

        # Construit la grille 2D : [lignes+1][colonnes+1] avec en-tetes
        n_cols = len(xs) + 1
        n_lignes = len(ys) + 1
        grille = [[""] * n_cols for _ in range(n_lignes)]
        # En-tete ligne 0 : lib du groupe + libs colonnes X
        grille[0][0] = gr.get("lib_groupe") or ""
        for ic, x in enumerate(xs, start=1):
            grille[0][ic] = x.get("lib") or ""
        # En-tetes colonne 0 (libs Y)
        for il, y in enumerate(ys, start=1):
            grille[il][0] = y.get("lib") or ""
        # Cellules
        for il, y in enumerate(ys, start=1):
            for ic, x in enumerate(xs, start=1):
                mnt = cell_map.get((int(x["id_groupe_rem_x"]), int(y["id_groupe_rem_y"])), 0)
                grille[il][ic] = f"{mnt:.2f} €".replace(".", ",") if mnt else ""

        out.append({
            "id_groupe_rem": id_gr,
            "lib_groupe": gr.get("lib_groupe") or "",
            "grille": grille,
        })
    return out


def _replace_all(doc, tag_dict: dict[str, str]) -> None:
    """Remplace tous les tags {tag} dans le document (paragraphes,
    tableaux, en-tetes, pieds de page). Preserve la mise en forme
    des runs quand possible."""
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for tag, val in tag_dict.items():
                if tag not in p.text: continue
                # Recomposition du texte : on concatene tous les runs,
                # on remplace, puis on remet dans le 1er run et on vide
                # les autres (perd les formatages intermediaires mais
                # marche pour le 90% des cas).
                full = "".join(r.text for r in p.runs)
                if tag not in full: continue
                new_full = full.replace(tag, val or "")
                if p.runs:
                    p.runs[0].text = new_full
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = new_full

    replace_in_paragraphs(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
                for inner_tbl in cell.tables:
                    for inner_row in inner_tbl.rows:
                        for inner_cell in inner_row.cells:
                            replace_in_paragraphs(inner_cell.paragraphs)
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)


def _insert_tableau_apres_paragraph(doc, para, grille: list[list[str]]) -> None:
    """Insere un tableau python-docx apres le paragraphe donne."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    n_lignes = len(grille)
    n_cols = len(grille[0]) if grille else 0
    if n_lignes == 0 or n_cols == 0:
        return

    tbl = doc.add_table(rows=n_lignes, cols=n_cols)
    tbl.style = "Table Grid"
    for il, row_data in enumerate(grille):
        for ic, val in enumerate(row_data):
            cell = tbl.rows[il].cells[ic]
            cell.text = str(val or "")
            # Formatage : gras + couleur pour ligne 0 et colonne 0
            if il == 0 or ic == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)

    # Deplace le tableau juste apres le paragraphe cible
    tbl_element = tbl._element
    para._element.addnext(tbl_element)


def _insert_tableaux_rem(doc, groupes_rem: list[dict]) -> None:
    """Insere un tableau REM apres CHAQUE tag TABLEAU_REM."""
    for grp in groupes_rem:
        # Cherche un paragraphe contenant TABLEAU_REM
        target = None
        for p in doc.paragraphs:
            if "TABLEAU_REM" in p.text:
                target = p
                break
        if target is None: continue
        # Insere le tableau apres le paragraphe
        _insert_tableau_apres_paragraph(doc, target, grp["grille"])
        # Remplace le tag par le libGroupe (cf WinDev :
        # FragmentDeb.Texte = "TABLEAU_REM"+RC+LibGroupe)
        for r in target.runs:
            r.text = ""
        if target.runs:
            target.runs[0].text = grp["lib_groupe"]
        else:
            target.text = grp["lib_groupe"]

    # Nettoie les tags TABLEAU_REM restants
    for p in doc.paragraphs:
        if "TABLEAU_REM" in p.text:
            for r in p.runs:
                r.text = r.text.replace("TABLEAU_REM", "")


def generer_contrat_docx(
    id_doc_courtage: int, id_distrib: int, id_gerant: int,
    secteur: str, date_signature: str,
    date_avenant: str = "",
) -> bytes:
    """Genere le DOCX rempli. Reproduit btn 'Ticket Omaya' WinDev
    (partie publipostage seulement - pas de PDF/FTP/ticket ici)."""
    from docx import Document

    # 1. Charge le template DOCX depuis pgt_doc_courtage
    db = get_pg_connection("rh")
    doc_row = db.query_one(
        """SELECT titre, contenu, id_groupe_operateur, id_ste
             FROM rh.pgt_doc_courtage
            WHERE id_doc_courtage = ? LIMIT 1""",
        (int(id_doc_courtage),),
    )
    if not doc_row or not doc_row.get("contenu"):
        raise ValueError("Document introuvable ou contenu vide")

    contenu = doc_row["contenu"]
    if isinstance(contenu, memoryview): contenu = bytes(contenu)
    id_groupe_op = int(doc_row.get("id_groupe_operateur") or 0)
    id_ste = int(doc_row.get("id_ste") or 0)
    titre = doc_row.get("titre") or ""

    doc = Document(io.BytesIO(contenu))

    # 2. Tags Distrib + Gerant salarie
    tags = _get_distrib_info(id_distrib)
    tags["SECTEUR_DISTRIB"] = secteur or ""
    tags["DATE_CTS"] = _fmt_date_fr(date_signature)
    tags["DOCTITRE"] = titre

    # 3. Tags STE editrice (sauf groupe 'Autres')
    if id_groupe_op != GROUPE_OP_AUTRE and id_ste:
        tags.update(_get_ste_info(id_ste))
    else:
        # Groupe 'Autres' : vide les tags STE pour eviter les {STE_...} residuels
        for k in ("STE_RS", "STE_APE", "STE_RCS", "STE_CAPITAL", "STE_ADR",
                  "STE_VILLE", "STE_GERANT_NOM", "STE_GERANT_TYPE",
                  "STE_SIRET", "STE_SIREN"):
            tags.setdefault(k, "")

    # 4. Tag DATE_AVENANT si le doc est un avenant
    if date_avenant and "AVENANT" in titre.upper():
        tags["DATE_AVENANT"] = _fmt_date_fr(date_avenant)

    # 5. Tags STE_LOGO / GER_SIGN / STE_CACHET / S_SIGN_DISTRIB / S_MENTION_DISTRIB
    # -> vides pour l'instant (images pas encore implementees)
    for img_tag in ("STE_LOGO", "GER_SIGN", "STE_CACHET",
                     "S_SIGN_DISTRIB", "S_MENTION_DISTRIB"):
        tags.setdefault(img_tag, "")

    # 6. Insertion des tableaux REM (sauf groupe 'Autres')
    if id_groupe_op != GROUPE_OP_AUTRE:
        groupes_rem = _get_groupes_rem_publipostage(id_distrib, id_groupe_op)
        _insert_tableaux_rem(doc, groupes_rem)
    else:
        # Groupe 'Autres' : juste nettoyer le tag TABLEAU_REM
        for p in doc.paragraphs:
            if "TABLEAU_REM" in p.text:
                for r in p.runs:
                    r.text = r.text.replace("TABLEAU_REM", "")

    # 7. Remplacements de tous les tags
    _replace_all(doc, tags)

    # 8. Sauve en bytes
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()


def _new_id_courtage() -> int:
    """ID entier 8 octets."""
    return int(datetime.now().strftime("%Y%m%d%H%M%S%f")[:17])


def create_societe_doc_courtage(
    id_distrib: int, id_gerant: int, id_doc_courtage: int,
    id_groupe_operateur: int, secteur: str, op_id: int,
) -> int:
    """Cree une entree societe_docCourtage cf WinDev 'Suivi d'edition'."""
    db = get_pg_connection("rh")
    id_new = _new_id_courtage()
    db.query(
        """INSERT INTO rh.pgt_societe_doc_courtage
              (id_societe_doc_courtage, id_salarie, id_distrib,
               id_doc_courtage, id_groupe_operateur, secteur,
               date_edition, recu, modif_date, modif_op, modif_elem)
           VALUES (?, ?, ?, ?, ?, ?, NOW(), FALSE, NOW(), ?, 'new')""",
        (id_new, int(id_gerant), int(id_distrib),
         int(id_doc_courtage), int(id_groupe_operateur),
         secteur or "", int(op_id)),
    )
    return id_new
