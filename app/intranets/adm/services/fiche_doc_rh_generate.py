"""
Generation d'un contrat de travail (transposition Fen_SalarieDocRH WinDev).

Bouton 'Ticket Omaya' :
  1. Charger le modele DOCX (pgt_doc_rh.contenu, bytea).
  2. Publipostage_Salarie : remplace S_TITRE, S_NOM, S_PRENOM, S_LNAISS,
     S_DEPNAISS, S_NUMSS, S_DNAISS, S_ADRESSE, S_CP, S_VILLE, S_GSM,
     FIN_PER_ESSAI, DATE_CTS, DATE_ANC, SECTEURAGENCE + cas AVENANT
     (DATE_AVENANT, DATE_AVENANT_FINESSAI).
  3. Publipostage_STE : remplace DOCTITRE, STE_RS, STE_APE, STE_RCS,
     STE_CAPITAL, STE_ADR, STE_VILLE, STE_SIREN, STE_SIRET, STE_GERANT_NOM,
     STE_GERANT_TYPE et inserre les images STE_LOGO (guimmick), GER_SIGN
     (gerant_signature), STE_CACHET (cachet_cial), + paraphe en pied de
     page. WinDev clone le logo+paraphe a chaque 'Art.', 'Annexe',
     'Mention', 'S_SIGN', '>' -- on conserve ce comportement.
  4. Cree 3 records :
       a. pgt_salarie_doc_rh (suivi d'edition, ModifElem='new')
       b. pgt_tk_demande_ctt_w (contenu = DOCX bytes, contrat_genere=1)
       c. pgt_tk_liste (ticket type 4, statut 1, service 'RH',
          destinataire = idDA)
  5. Convertit DOCX -> PDF via LibreOffice headless.
  6. Upload FTP du PDF dans /TempCttw/<idTicket>-cttW.pdf (servi par IIS
     interne sur https://interne.omaya.fr/TempCttw/).
"""

from __future__ import annotations

import io
import os
import subprocess
import tempfile
from datetime import datetime
from typing import Any

from app.core.config import SOFFICE_BIN
from app.core.database import get_connection
from app.core.database.pg import get_pg_connection
from app.shared.tickets.forms.cttw_pdf import ftp_upload


# --- Helpers --------------------------------------------------------------

def _str(v: Any) -> str:
    return "" if v is None else str(v)


def _int(v: Any) -> int:
    if v is None or v == "":
        return 0
    try:
        return int(v)
    except (TypeError, ValueError):
        return 0


def _new_id() -> int:
    """ID 8 octets timestamp (cf. WinDev idEntierDateHeureSys)."""
    n = datetime.now()
    return int(n.strftime("%Y%m%d%H%M%S") + f"{n.microsecond // 1000:03d}")


def _fmt_fr_date(v: Any) -> str:
    """Date / datetime / string -> 'DD/MM/YYYY' (WinDev maskDateSysteme)."""
    if v is None or v == "":
        return ""
    if isinstance(v, datetime):
        return v.strftime("%d/%m/%Y")
    s = str(v)
    if len(s) >= 10 and s[4] == "-" and s[7] == "-":
        return f"{s[8:10]}/{s[5:7]}/{s[0:4]}"
    return s


def _capitalize_first(s: str) -> str:
    if not s:
        return ""
    return s[0].upper() + s[1:].lower()


# --- Publipostage texte ---------------------------------------------------

def _iter_paragraphs(doc):
    """Itere sur tous les paragraphes (corps + tableaux + footers/headers)."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if footer is None:
                continue
            for p in footer.paragraphs:
                yield p
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if header is None:
                continue
            for p in header.paragraphs:
                yield p


def _replace_text(doc, mapping: dict[str, str]) -> None:
    """Remplace les tokens texte dans tous les runs (preserve formatting).

    Strategie : pour chaque paragraphe, concatener les runs s'ils
    contiennent des fragments d'un meme token (Word peut splitter
    'S_NOM' en plusieurs runs si on a edite le doc).

    Trie les cles par longueur decroissante pour eviter que 'DATE_AVENANT'
    soit substitue avant 'DATE_AVENANT_FINESSAI' (et casse ce dernier).
    """
    # Cles triees par longueur desc : les plus longues d'abord.
    sorted_items = sorted(mapping.items(), key=lambda kv: -len(kv[0]))
    for p in _iter_paragraphs(doc):
        if not any(tok in p.text for tok, _ in sorted_items):
            continue
        full = p.text
        new = full
        for tok, val in sorted_items:
            if tok in new:
                new = new.replace(tok, val or "")
        if new == full:
            continue
        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""


def _replace_token_with_image(doc, token: str, img_bytes: bytes, width_mm: float = 20) -> int:
    """Remplace chaque occurrence de `token` par l'image. Retourne le nombre
    de remplacements effectues. L'image est dimensionnee par largeur."""
    from docx.shared import Mm

    count = 0
    for p in _iter_paragraphs(doc):
        if token not in p.text:
            continue
        for run in list(p.runs):
            if token not in run.text:
                continue
            before, _, after = run.text.partition(token)
            run.text = before
            try:
                run.add_picture(io.BytesIO(img_bytes), width=Mm(width_mm))
                count += 1
            except Exception:
                pass
            if after:
                tail = p.add_run(after)
                tail.font.size = run.font.size
    return count


def _insert_image_before_tokens(
    doc, tokens: list[str], img_bytes: bytes, width_mm: float = 10
) -> int:
    """Insere l'image AVANT chaque occurrence de tokens. Retourne le nombre
    d'insertions effectuees. Reproduit le comportement WinDev qui clone
    le logo+paraphe a chaque 'Art.', 'Annexe', 'Mention', 'S_SIGN', '>'.

    NB : pour eviter les explosions de logos (cf. WinDev qui en met
    beaucoup), on plafonne a 100 par token."""
    from docx.shared import Mm

    count = 0
    max_per_token = 100
    for token in tokens:
        per_tok = 0
        for p in _iter_paragraphs(doc):
            if token not in p.text or per_tok >= max_per_token:
                continue
            for run in list(p.runs):
                if token not in run.text or per_tok >= max_per_token:
                    continue
                before, _, after = run.text.partition(token)
                run.text = before
                # Insere l'image avant le token, puis retape le token + after
                try:
                    run.add_picture(io.BytesIO(img_bytes), width=Mm(width_mm))
                    per_tok += 1
                    count += 1
                except Exception:
                    pass
                tail = p.add_run(token + after)
                tail.font.size = run.font.size
    return count


def _add_paraphe_footer(doc, img_bytes: bytes, width_mm: float = 12) -> None:
    """Ajoute le paraphe gerant dans le pied de page (aligne a droite).

    Copie de cttw_pdf.py._add_paraphe_footer (sections liees au precedent
    partagent le meme footer XML)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    seen: set[int] = set()
    for section in doc.sections:
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                if footer is None or footer.is_linked_to_previous:
                    continue
                key = id(footer._element)
            except Exception:
                continue
            if key in seen:
                continue
            seen.add(key)
            para = footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                para.add_run().add_picture(io.BytesIO(img_bytes), width=Mm(width_mm))
            except Exception:
                pass


# --- Conversion DOCX -> PDF ----------------------------------------------

def _docx_to_pdf(docx_bytes: bytes) -> bytes:
    """LibreOffice headless docx -> pdf. Retourne les bytes PDF."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        # soffice --headless --convert-to pdf --outdir tmpdir input.docx
        try:
            subprocess.run(
                [
                    SOFFICE_BIN, "--headless", "--norestore", "--nologo",
                    "--nolockcheck", "--convert-to", "pdf", "--outdir",
                    tmpdir, docx_path,
                ],
                check=True,
                capture_output=True,
                timeout=120,
                creationflags=0x08000000,  # CREATE_NO_WINDOW
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"soffice : {e.stderr.decode('utf-8', 'ignore')}")
        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("soffice n'a pas produit de PDF")
        with open(pdf_path, "rb") as f:
            return f.read()


# --- Donnees salarie / societe -------------------------------------------

def _load_salarie(id_salarie: int) -> dict:
    """Charge salarie + embauche + coordonnees + organigramme courant."""
    db = get_pg_connection("rh")
    sal = db.query_one(
        """SELECT id_salarie, civilite, nom, prenom, num_ss,
                  date_naiss, lieu_naiss, dep_naiss
           FROM rh.pgt_salarie WHERE id_salarie = ?""",
        (int(id_salarie),),
    ) or {}
    coord = db.query_one(
        """SELECT adresse1, cp, ville, tel_mob
           FROM rh.pgt_salarie_coordonnees WHERE id_salarie = ?""",
        (int(id_salarie),),
    ) or {}
    emb = db.query_one(
        """SELECT id_ste, date_debut, date_fin_per_essai, date_anciennete
           FROM rh.pgt_salarie_embauche WHERE id_salarie = ?""",
        (int(id_salarie),),
    ) or {}
    orga = db.query_one(
        """SELECT o.idorganigramme, o.secteur, o.id_type_produit
           FROM rh.pgt_salarie_organigramme so
           LEFT JOIN rh.pgt_organigramme o ON o.idorganigramme = so.idorganigramme
           WHERE so.id_salarie = ?
             AND so.modif_elem NOT LIKE '%suppr%'
             AND COALESCE(so.aff_actif, FALSE) = TRUE
           ORDER BY so.date_debut DESC NULLS LAST
           LIMIT 1""",
        (int(id_salarie),),
    ) or {}
    return {"salarie": sal, "coord": coord, "embauche": emb, "orga": orga}


def _load_societe(id_ste: int) -> dict:
    """Charge la societe avec ses images binaires."""
    if not id_ste:
        return {}
    db = get_pg_connection("rh")
    row = db.query_one(
        """SELECT id_ste, raison_sociale, code_ape, rcs, capital,
                  adresse1, cp, ville, siren, siret,
                  gerant_nom, gerant_type,
                  guimmick, gerant_paraphe, gerant_signature, cachet_cial
           FROM rh.pgt_societe WHERE id_ste = ?""",
        (int(id_ste),),
    )
    return row or {}


def _bytes_or_none(v: Any) -> bytes | None:
    if v is None:
        return None
    if isinstance(v, (bytes, bytearray)):
        return bytes(v)
    if isinstance(v, memoryview):
        return v.tobytes()
    return None


# --- Generation : publipostage DOCX (etapes 1-8) -------------------------

def _build_publiposted_docx(
    *,
    id_salarie: int,
    id_doc_rh: int,
    date_avenant: str = "",
) -> dict:
    """Charge le modele (DOCX ou HTML) et fait le publipostage.

    Retourne :
    - DOCX : {docx_bytes, is_docx=True, titre_doc, id_type_doc, id_da,
             idorganigramme, nom_salarie}
    - HTML : {html_bytes, is_docx=False, ...meme cles...}
    Pas d'ecriture DB ni de conversion PDF.
    """
    from app.intranets.adm.services.ctt_travail import is_docx as _is_docx
    from docx import Document  # noqa: PLC0415 -- lazy

    # 1. Charger le modele
    db_rh = get_pg_connection("rh")
    model = db_rh.query_one(
        """SELECT id_doc_rh, titre, id_type_doc, contenu, id_ste,
                  id_type_produit
           FROM rh.pgt_doc_rh WHERE id_doc_rh = ?""",
        (int(id_doc_rh),),
    )
    if not model:
        raise ValueError(f"Modele doc RH introuvable (id={id_doc_rh})")
    raw = _bytes_or_none(model.get("contenu"))
    if not raw:
        raise ValueError("Le modele de doc RH n'a pas de contenu.")
    titre_doc = _str(model.get("titre"))
    id_type_doc = _int(model.get("id_type_doc"))
    is_docx_format = _is_docx(raw)
    docx_raw = raw if is_docx_format else b""

    # 2. Donnees salarie + societe
    data = _load_salarie(int(id_salarie))
    sal = data["salarie"]
    coord = data["coord"]
    emb = data["embauche"]
    orga = data["orga"]
    id_ste = _int(emb.get("id_ste"))
    societe = _load_societe(id_ste)

    # 3. Determiner idDA (1er DA/DR de l'organigramme parent/courant)
    id_da = _find_id_da(int(id_salarie), _int(orga.get("idorganigramme")))

    # 4. Charger le DOCX en memoire (si format DOCX uniquement)
    doc = Document(io.BytesIO(docx_raw)) if is_docx_format else None

    # 5. Publipostage_Salarie
    civilite = _int(sal.get("civilite"))
    nom = _str(sal.get("nom"))
    prenom = _str(sal.get("prenom"))
    is_avenant = "AVENANT" in titre_doc.upper()
    txt_map = {
        "S_TITRE": "Mr." if civilite == 1 else "Mme",
        "S_NOM": nom,
        "S_PRENOM": _capitalize_first(prenom),
        "S_LNAISS": _str(sal.get("lieu_naiss")),
        "S_DEPNAISS": _str(sal.get("dep_naiss")),
        "S_NUMSS": _str(sal.get("num_ss")),
        "S_DNAISS": _fmt_fr_date(sal.get("date_naiss")),
        "S_ADRESSE": _str(coord.get("adresse1")),
        "S_CP": _str(coord.get("cp")),
        "S_VILLE": _str(coord.get("ville")),
        "S_GSM": _str(coord.get("tel_mob")),
        "FIN_PER_ESSAI": _fmt_fr_date(emb.get("date_fin_per_essai")),
        "DATE_CTS": _fmt_fr_date(emb.get("date_debut")),
        "DATE_ANC": _fmt_fr_date(emb.get("date_anciennete")),
        "SECTEURAGENCE": _str(orga.get("secteur")),
    }
    if is_avenant and date_avenant:
        from datetime import date as _date
        try:
            d = datetime.fromisoformat(date_avenant).date()
            # DateAvPerEssai = DateAv + 3 mois - 1 jour (cf. WinDev)
            mois = d.month + 3
            annee = d.year
            while mois > 12:
                mois -= 12
                annee += 1
            from calendar import monthrange
            jmax = monthrange(annee, mois)[1]
            jour = min(d.day, jmax)
            d_fin = _date(annee, mois, jour)
            # -1 jour
            from datetime import timedelta
            d_fin = d_fin - timedelta(days=1)
            txt_map["DATE_AVENANT"] = d.strftime("%d/%m/%Y")
            txt_map["DATE_AVENANT_FINESSAI"] = d_fin.strftime("%d/%m/%Y")
        except Exception:
            pass

    # 6. Publipostage_STE (texte)
    raison_sociale = _str(societe.get("raison_sociale"))
    txt_map.update({
        "DOCTITRE": titre_doc,
        "STE_RS": raison_sociale,
        "STE_APE": _str(societe.get("code_ape")),
        "STE_RCS": _str(societe.get("rcs")),
        "STE_CAPITAL": _str(societe.get("capital")),
        "STE_ADR": " ".join(
            x for x in (
                _str(societe.get("adresse1")),
                _str(societe.get("cp")),
                _str(societe.get("ville")),
            ) if x
        ),
        "STE_VILLE": _str(societe.get("ville")),
        "STE_SIREN": _str(societe.get("siren")),
        "STE_SIRET": _str(societe.get("siret")),
        "STE_GERANT_NOM": _str(societe.get("gerant_nom")),
        "STE_GERANT_TYPE": _str(societe.get("gerant_type")),
    })
    # Branche DOCX vs HTML
    if doc is not None:
        _replace_text(doc, txt_map)

        # 7. Insertions images (DOCX uniquement)
        img_guimmick = _bytes_or_none(societe.get("guimmick"))
        img_paraphe = _bytes_or_none(societe.get("gerant_paraphe"))
        img_signature = _bytes_or_none(societe.get("gerant_signature"))
        img_cachet = _bytes_or_none(societe.get("cachet_cial"))

        if img_guimmick:
            _replace_token_with_image(doc, "STE_LOGO", img_guimmick, width_mm=20)
        if img_signature:
            _replace_token_with_image(doc, "GER_SIGN", img_signature, width_mm=40)
        if img_cachet:
            _replace_token_with_image(doc, "STE_CACHET", img_cachet, width_mm=30)
        if img_paraphe:
            _add_paraphe_footer(doc, img_paraphe, width_mm=12)

        # 8. Sauver le DOCX
        docx_out = io.BytesIO()
        doc.save(docx_out)
        content_bytes = docx_out.getvalue()
    else:
        # Format HTML : substitution texte directe (les images sont gerees
        # par generer_pdf_publiposte au moment du rendu PDF).
        # Tri par longueur desc -> DATE_AVENANT_FINESSAI traite avant
        # DATE_AVENANT.
        body_html = raw.decode("utf-8", errors="ignore")
        for k, v in sorted(txt_map.items(), key=lambda kv: -len(kv[0])):
            if k == "STE_LOGO":  # gere par generer_pdf_publiposte
                continue
            body_html = body_html.replace(k, str(v))
        content_bytes = body_html.encode("utf-8")

    return {
        "docx_bytes": content_bytes,  # nom historique (bytea stocke)
        "is_docx": is_docx_format,
        "titre_doc": titre_doc,
        "id_type_doc": id_type_doc,
        "id_da": id_da,
        "idorganigramme": _int(orga.get("idorganigramme")),
        "nom_salarie": f"{nom} {_capitalize_first(prenom)}".strip(),
    }


# --- Generation complete -------------------------------------------------

def generate_cttw(
    *,
    id_salarie: int,
    id_doc_rh: int,
    op_id: int,
    date_avenant: str = "",
) -> dict:
    """Genere le contrat de travail (DOCX + PDF) + cree les 3 records.

    Retourne {ok, id_ticket, id_tk_demande_ctt_w, id_salarie_doc_rh,
              pdf_url, id_da, type_doc_lib}.
    """
    built = _build_publiposted_docx(
        id_salarie=id_salarie, id_doc_rh=id_doc_rh, date_avenant=date_avenant
    )
    docx_bytes = built["docx_bytes"]
    titre_doc = built["titre_doc"]
    id_type_doc = built["id_type_doc"]
    id_da = built["id_da"]
    idorganigramme = built["idorganigramme"]
    db_rh = get_pg_connection("rh")

    # 9. Creer pgt_salarie_doc_rh (suivi d'edition)
    id_salarie_doc_rh = _new_id()
    db_rh.query(
        """INSERT INTO rh.pgt_salarie_doc_rh
              (id_salarie_doc_rh, id_doc_rhtype, id_salarie, id_da,
               date_edition, recu,
               modif_date, modif_op, modif_elem)
           VALUES (?, ?, ?, ?, NOW(), FALSE, NOW(), ?, 'new')""",
        (
            id_salarie_doc_rh,
            int(id_doc_rh),
            int(id_salarie),
            id_da,
            int(op_id),
        ),
    )

    # 10. Creer pgt_tk_demande_ctt_w (contenu = DOCX bytes)
    id_demande = _new_id()
    id_ticket = _new_id()
    db_trh = get_pg_connection("ticket_rh")
    db_trh.query(
        """INSERT INTO ticket_rh.pgt_tk_demande_ctt_w
              (id_demande_contrat_w, id_doc_rhedit, idorganigramme,
               id_salarie, id_da, type_ctt_w, titre_contrat, contenu,
               contrat_genere, contrat_valide, contrat_signe,
               contrat_annul, id_tk_liste,
               modif_date, modif_op, modif_elem)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, TRUE, FALSE, FALSE, FALSE, ?,
                   NOW(), ?, 'new')""",
        (
            id_demande,
            id_salarie_doc_rh,
            idorganigramme,
            int(id_salarie),
            id_da,
            str(id_type_doc),
            titre_doc,
            docx_bytes,
            id_ticket,
            int(op_id),
        ),
    )

    # 11. Creer pgt_tk_liste (ticket type 4 = demande contrat W, service RH)
    db_t = get_pg_connection("ticket")
    db_t.query(
        """INSERT INTO ticket.pgt_tk_liste
              (id_tk_liste, date_crea, op_crea, op_dest, service,
               id_tk_type_demande, id_tk_statut, cloturee,
               modif_date, modif_op, modif_elem,
               op_traitement_staff, ordre_traitement_staff)
           VALUES (?, NOW(), ?, ?, 'RH', 4, 1, FALSE, NOW(), ?, 'new',
                   0, 0)""",
        (id_ticket, int(op_id), id_da, int(op_id)),
    )

    # 12. Generation PDF (WeasyPrint + footer + SAUTDEPAGE)
    from app.intranets.adm.services.ctt_travail import (
        _docx_to_html,
        generer_pdf_publiposte,
    )
    if built.get("is_docx", True):
        # DOCX publiposte -> HTML via mammoth (images embedded en data:)
        body_html_for_pdf = _docx_to_html(docx_bytes)
    else:
        # Contenu deja HTML publiposte
        body_html_for_pdf = docx_bytes.decode("utf-8", errors="ignore")
    id_ste_for_pdf = _int(
        (_load_salarie(int(id_salarie)).get("embauche") or {}).get("id_ste")
    )
    pdf_bytes = generer_pdf_publiposte(
        body_html_for_pdf, id_ste=id_ste_for_pdf,
    )
    if not pdf_bytes:
        raise RuntimeError("Generation PDF (WeasyPrint) a echoue.")

    # 13. Upload FTP TempCttw/ (servi par IIS interne sous interne.omaya.fr/TempCttw/)
    pdf_name = f"{id_ticket}-cttW.pdf"
    ftp_path = os.getenv("FTP_TEMPCTTW_PATH", "/OMAYA/TempCttw")
    try:
        ftp_upload(ftp_path, pdf_name, pdf_bytes)
    except Exception:
        # On ne fait pas echouer la generation pour un soucis FTP : le
        # PDF est deja en base via pgt_tk_demande_ctt_w.contenu.
        import sys, traceback
        traceback.print_exc(file=sys.stderr)

    return {
        "ok": True,
        "id_ticket": str(id_ticket),
        "id_tk_demande_ctt_w": str(id_demande),
        "id_salarie_doc_rh": str(id_salarie_doc_rh),
        "id_da": str(id_da),
        "type_doc_lib": titre_doc,
        "pdf_url": f"https://interne.omaya.fr/TempCttw/{pdf_name}",
    }


def preview_cttw_pdf(
    *,
    id_salarie: int,
    id_doc_rh: int,
    date_avenant: str = "",
) -> dict:
    """Genere le PDF pour aperçu/export, SANS rien ecrire en base.

    Bouton 'Export PDF' Fen_SalarieDocRH. Supporte 2 formats de modele :
    - DOCX (legacy) : publipostage via python-docx + mammoth -> HTML.
    - HTML (nouveau Fen_EditionDocRH) : substitution texte directe.
    Puis WeasyPrint avec footer auto + SAUTDEPAGE + conversion <font>.
    """
    from app.intranets.adm.services.ctt_travail import (
        _docx_to_html,
        generer_pdf_publiposte,
        is_docx,
    )

    # 1. Charge le modele pour detecter le format (DOCX ou HTML).
    db_rh = get_pg_connection("rh")
    model = db_rh.query_one(
        """SELECT id_doc_rh, titre, contenu
           FROM rh.pgt_doc_rh WHERE id_doc_rh = ?""",
        (int(id_doc_rh),),
    )
    if not model:
        raise ValueError(f"Modele doc RH introuvable (id={id_doc_rh}).")
    content = _bytes_or_none(model.get("contenu"))
    if not content:
        raise ValueError("Le modele de doc RH n'a pas de contenu.")
    titre_doc = _str(model.get("titre"))

    # 2. Donnees salarie + societe pour le footer + nom de fichier.
    data = _load_salarie(int(id_salarie))
    id_ste = _int((data.get("embauche") or {}).get("id_ste"))
    sal = data.get("salarie") or {}
    nom = _str(sal.get("nom"))
    prenom = _str(sal.get("prenom"))
    nom_salarie = f"{nom} {_capitalize_first(prenom)}".strip()

    # 3. Selon le format, prepare le HTML publiposte.
    if is_docx(content):
        # Format DOCX : reuse _build_publiposted_docx (variables + images
        # python-docx) puis mammoth -> HTML (images data: URLs).
        built = _build_publiposted_docx(
            id_salarie=id_salarie,
            id_doc_rh=id_doc_rh,
            date_avenant=date_avenant,
        )
        body_html = _docx_to_html(built["docx_bytes"])
    else:
        # Format HTML : substitution texte directe (les images STE_LOGO
        # / GER_SIGN / STE_CACHET sont substituees par
        # generer_pdf_publiposte via extra_images).
        body_html = _build_html_publiposte(
            content.decode("utf-8", errors="ignore"),
            id_salarie=id_salarie,
            titre_doc=titre_doc,
            date_avenant=date_avenant,
        )

    pdf_bytes = generer_pdf_publiposte(body_html, id_ste=id_ste)
    if not pdf_bytes:
        raise RuntimeError("Generation PDF (WeasyPrint) a echoue.")

    safe_nom = nom_salarie.replace(" ", "_") or str(id_salarie)
    safe_titre = titre_doc.replace("/", "-").replace("\\", "-")[:80]
    filename = (
        f"{safe_nom}_{safe_titre}.pdf" if safe_titre else f"{safe_nom}.pdf"
    )
    return {"pdf_bytes": pdf_bytes, "filename": filename}


def _build_html_publiposte(
    body_html: str,
    *,
    id_salarie: int,
    titre_doc: str = "",
    date_avenant: str = "",
) -> str:
    """Substitue les variables S_* / STE_* / DATE_* dans un contenu HTML
    (Fen_EditionDocRH). Retourne le HTML publiposte (sans toucher aux
    images : STE_LOGO/GER_SIGN/STE_CACHET sont gerees par
    generer_pdf_publiposte)."""
    data = _load_salarie(int(id_salarie))
    sal = data.get("salarie") or {}
    coord = data.get("coord") or {}
    emb = data.get("embauche") or {}
    orga = data.get("orga") or {}
    id_ste = _int(emb.get("id_ste"))
    societe = _load_societe(id_ste) if id_ste else {}

    civilite = _int(sal.get("civilite"))
    nom = _str(sal.get("nom"))
    prenom = _str(sal.get("prenom"))
    is_avenant = "AVENANT" in (titre_doc or "").upper()
    txt_map = {
        "S_TITRE": "Mr." if civilite == 1 else "Mme",
        "S_NOM": nom,
        "S_PRENOM": _capitalize_first(prenom),
        "S_LNAISS": _str(sal.get("lieu_naiss")),
        "S_DEPNAISS": _str(sal.get("dep_naiss")),
        "S_NUMSS": _str(sal.get("num_ss")),
        "S_DNAISS": _fmt_fr_date(sal.get("date_naiss")),
        "S_ADRESSE": _str(coord.get("adresse1")),
        "S_CP": _str(coord.get("cp")),
        "S_VILLE": _str(coord.get("ville")),
        "S_GSM": _str(coord.get("tel_mob")),
        "FIN_PER_ESSAI": _fmt_fr_date(emb.get("date_fin_per_essai")),
        "DATE_CTS": _fmt_fr_date(emb.get("date_debut")),
        "DATE_ANC": _fmt_fr_date(emb.get("date_anciennete")),
        "SECTEURAGENCE": _str(orga.get("secteur")),
        "DOCTITRE": titre_doc,
        "STE_RS": _str(societe.get("raison_sociale")),
        "STE_APE": _str(societe.get("code_ape")),
        "STE_RCS": _str(societe.get("rcs")),
        "STE_CAPITAL": _str(societe.get("capital")),
        "STE_ADR": " ".join(
            x for x in (
                _str(societe.get("adresse1")),
                _str(societe.get("cp")),
                _str(societe.get("ville")),
            ) if x
        ),
        "STE_VILLE": _str(societe.get("ville")),
        "STE_SIREN": _str(societe.get("siren")),
        "STE_SIRET": _str(societe.get("siret")),
        "STE_GERANT_NOM": _str(societe.get("gerant_nom")),
        "STE_GERANT_TYPE": _str(societe.get("gerant_type")),
    }
    if is_avenant and date_avenant:
        from datetime import date as _date
        try:
            d = datetime.fromisoformat(date_avenant).date()
            mois = d.month + 3
            annee = d.year
            while mois > 12:
                mois -= 12
                annee += 1
            from calendar import monthrange
            jmax = monthrange(annee, mois)[1]
            jour = min(d.day, jmax)
            d_fin = _date(annee, mois, jour)
            from datetime import timedelta
            d_fin = d_fin - timedelta(days=1)
            txt_map["DATE_AVENANT"] = d.strftime("%d/%m/%Y")
            txt_map["DATE_AVENANT_FINESSAI"] = d_fin.strftime("%d/%m/%Y")
        except Exception:
            pass

    # Tri par longueur desc : DATE_AVENANT_FINESSAI avant DATE_AVENANT.
    for k, v in sorted(txt_map.items(), key=lambda kv: -len(kv[0])):
        if k == "STE_LOGO":  # gere par generer_pdf_publiposte
            continue
        body_html = body_html.replace(k, str(v))
    return body_html


def _find_id_da(id_salarie: int, id_organigramme: int) -> int:
    """Trouve le 1er DA/DR (responsable d'equipe) rattache au meme
    organigramme ou a l'organigramme parent."""
    if not id_organigramme:
        return 0
    db = get_pg_connection("rh")
    # Recherche : salaries actifs avec resp_equipe=true sur cet orga ou parent
    row = db.query_one(
        """SELECT se.id_salarie
           FROM rh.pgt_salarie_organigramme so
           LEFT JOIN rh.pgt_organigramme o ON o.idorganigramme = so.idorganigramme
           LEFT JOIN rh.pgt_salarie_embauche se ON se.id_salarie = so.id_salarie
           WHERE (so.idorganigramme = ?
                  OR so.idorganigramme = (SELECT id_parent FROM rh.pgt_organigramme
                                          WHERE idorganigramme = ?))
             AND COALESCE(so.aff_actif, FALSE) = TRUE
             AND so.modif_elem NOT LIKE '%suppr%'
             AND COALESCE(se.resp_equipe, FALSE) = TRUE
           ORDER BY so.date_debut DESC NULLS LAST
           LIMIT 1""",
        (int(id_organigramme), int(id_organigramme)),
    )
    return _int(row.get("id_salarie")) if row else 0
