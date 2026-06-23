"""
Service Fen_ListeDocRH (ADM, Salaries -> Liste des contrats de travail).

Transposition WinDev :
- Table_DocRH_Actif : JOIN doc_rh + doc_rhtype (type doc) + type_produit
  (logo + libelle) filtre par DocActif (glissiere haut droit).
- Boutons : Nouveau / Dupliquer / Supprimer / Modifier / Archiver.

L'edition d'un docRH (Fen_EditionDocRH) sera dans un module dedie.
"""

from __future__ import annotations

from datetime import datetime
from typing import Any

from app.core.database.pg import get_pg_connection
from app.shared.notifications.mail import envoi_mail


def _str(v: Any) -> str:
    return "" if v is None else str(v)


def _int(v: Any) -> int:
    if v is None or v == "":
        return 0
    try:
        return int(v)
    except (TypeError, ValueError):
        return 0


def _new_id() -> int:
    """idEntierDateHeureSys WinDev."""
    n = datetime.now()
    return int(n.strftime("%Y%m%d%H%M%S") + f"{n.microsecond // 1000:03d}")


def list_docs(doc_actif: bool = True) -> list[dict]:
    """Cf. requete ReqListeDocRH WinDev (filtre DocActif via glissiere)."""
    db = get_pg_connection("rh")
    rows = db.query(
        """SELECT d.id_doc_rh, d.id_type_doc, d.titre, d.info_cpl,
                  d.id_type_produit, d.datecrea, d.doc_actif,
                  d.prioritaire, d.modif_date,
                  d.id_ste, d.doc_dpae,
                  t.lib_type,
                  tp.lib AS lib_produit
             FROM rh.pgt_doc_rh d
        LEFT JOIN rh.pgt_doc_rhtype t ON t.id_type_doc = d.id_type_doc
        LEFT JOIN rh.pgt_type_produit tp ON tp.id_type_produit = d.id_type_produit
            WHERE (d.modif_elem IS NULL OR d.modif_elem NOT LIKE '%suppr%')
              AND COALESCE(d.doc_actif, FALSE) = ?
         ORDER BY tp.lib ASC NULLS LAST, d.titre ASC""",
        (bool(doc_actif),),
    ) or []
    return [
        {
            "id_doc_rh": str(r.get("id_doc_rh") or ""),
            "id_type_doc": str(_int(r.get("id_type_doc")) or ""),
            "lib_type": _str(r.get("lib_type")),
            "titre": _str(r.get("titre")),
            "info_cpl": _str(r.get("info_cpl")),
            "id_type_produit": str(_int(r.get("id_type_produit")) or ""),
            "lib_produit": _str(r.get("lib_produit")),
            "id_ste": str(_int(r.get("id_ste")) or ""),
            "doc_dpae": bool(r.get("doc_dpae")),
            "prioritaire": bool(r.get("prioritaire")),
            "datecrea": _str(r.get("datecrea"))[:19],
            "modif_date": _str(r.get("modif_date"))[:19],
        }
        for r in rows
    ]


def duplicate_doc(id_doc_rh: int, op_id: int, user_login: str = "",
                  user_prenom: str = "") -> dict:
    """Btn Dupliquer : copie le doc avec une nouvelle pk, force
    prioritaire=False et doc_dpae=False. Cf. WinDev."""
    db = get_pg_connection("rh")
    src = db.query_one(
        "SELECT * FROM rh.pgt_doc_rh WHERE id_doc_rh = ? LIMIT 1",
        (int(id_doc_rh),),
    )
    if not src:
        return {"ok": False, "error": "Document introuvable"}

    new_id = _new_id()
    db.query(
        """INSERT INTO rh.pgt_doc_rh
              (id_doc_rh, id_type_doc, titre, info_cpl, id_type_produit,
               contenu, datecrea, doc_actif, prioritaire, id_ste, doc_dpae,
               doc_dpae_distrib, id_tk_type_photo_dpae,
               modif_date, modif_op, modif_elem)
           VALUES (?, ?, ?, ?, ?,
                   ?, NOW(), TRUE, FALSE, ?, FALSE,
                   ?, ?,
                   NOW(), ?, 'new')""",
        (
            new_id,
            _int(src.get("id_type_doc")),
            _str(src.get("titre")),
            _str(src.get("info_cpl")),
            _int(src.get("id_type_produit")),
            src.get("contenu"),
            _int(src.get("id_ste")),
            bool(src.get("doc_dpae_distrib")),
            _int(src.get("id_tk_type_photo_dpae")),
            int(op_id),
        ),
    )

    # Mail a marie@exosphere.fr (cf. WinDev usersCial <> 256)
    if int(op_id) != 256 and user_login:
        try:
            envoi_mail(
                sujet=f"Ctt de travail duplique {_str(src.get('titre'))} - {_str(src.get('info_cpl'))}",
                html=(
                    f"<p>Bonjour,</p>"
                    f"<p>Le contrat de travail <b>{_str(src.get('titre'))} - "
                    f"{_str(src.get('info_cpl'))}</b> vient d'etre duplique par "
                    f"<b>{user_prenom or user_login}</b>.</p>"
                    f"<p>Cordialement.</p>"
                ),
                destinataires=["marie@exosphere.fr"],
                expediteur=user_login,
            )
        except Exception:
            pass
    return {"ok": True, "id_doc_rh": str(new_id)}


def archive_doc(id_doc_rh: int, op_id: int) -> dict:
    """Btn Archiver : doc_actif=FALSE."""
    db = get_pg_connection("rh")
    db.query(
        """UPDATE rh.pgt_doc_rh
              SET doc_actif = FALSE,
                  modif_date = NOW(),
                  modif_op = ?,
                  modif_elem = 'modif'
            WHERE id_doc_rh = ?""",
        (int(op_id), int(id_doc_rh)),
    )
    return {"ok": True}


def restore_doc(id_doc_rh: int, op_id: int) -> dict:
    """Re-actif depuis l'archive."""
    db = get_pg_connection("rh")
    db.query(
        """UPDATE rh.pgt_doc_rh
              SET doc_actif = TRUE,
                  modif_date = NOW(),
                  modif_op = ?,
                  modif_elem = 'modif'
            WHERE id_doc_rh = ?""",
        (int(op_id), int(id_doc_rh)),
    )
    return {"ok": True}


def delete_doc(id_doc_rh: int, op_id: int) -> dict:
    """Btn Supprimer : soft delete (modif_elem='suppr')."""
    db = get_pg_connection("rh")
    db.query(
        """UPDATE rh.pgt_doc_rh
              SET modif_date = NOW(),
                  modif_op = ?,
                  modif_elem = 'suppr'
            WHERE id_doc_rh = ?""",
        (int(op_id), int(id_doc_rh)),
    )
    return {"ok": True}


# ===========================================================================
# Fen_EditionDocRH (edition d'un doc RH)
# ===========================================================================


def list_types_doc() -> list[dict]:
    """Combo Type Doc."""
    db = get_pg_connection("rh")
    rows = db.query(
        """SELECT id_type_doc, lib_type FROM rh.pgt_doc_rhtype
            WHERE (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
         ORDER BY lib_type""",
    ) or []
    return [
        {"id_type_doc": str(_int(r.get("id_type_doc"))), "lib": _str(r.get("lib_type"))}
        for r in rows
    ]


def list_types_produit() -> list[dict]:
    """Combo Produit."""
    db = get_pg_connection("rh")
    rows = db.query(
        """SELECT id_type_produit, lib FROM rh.pgt_type_produit
            WHERE (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
         ORDER BY lib""",
    ) or []
    return [
        {"id_type_produit": str(_int(r.get("id_type_produit"))), "lib": _str(r.get("lib"))}
        for r in rows
    ]


def list_societes_actives() -> list[dict]:
    """Combo Societe (Fen_EditionDocRH combo ListeSTE + meta du doc)."""
    db = get_pg_connection("rh")
    rows = db.query(
        """SELECT id_ste, raison_sociale, rs_interne FROM rh.pgt_societe
            WHERE COALESCE(is_actif, FALSE) = TRUE
              AND (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
         ORDER BY raison_sociale""",
    ) or []
    return [
        {
            "id_ste": str(_int(r.get("id_ste"))),
            "raison_sociale": _str(r.get("raison_sociale")),
            "rs_interne": _str(r.get("rs_interne")),
        }
        for r in rows
    ]


def list_types_photo_dpae() -> list[dict]:
    """Combo Type Photo (TK_TypePhotoDPAE)."""
    db = get_pg_connection("ticket_dpae")
    rows = db.query(
        """SELECT id_tk_type_photo_dpae, lib_type_doc
             FROM ticket_dpae.pgt_tk_type_photo_dpae
            WHERE COALESCE(desactiver, FALSE) = FALSE
              AND (modif_elem IS NULL OR modif_elem NOT LIKE '%suppr%')
         ORDER BY lib_type_doc""",
    ) or []
    return [
        {
            "id_tk_type_photo_dpae": str(_int(r.get("id_tk_type_photo_dpae"))),
            "lib": _str(r.get("lib_type_doc")),
        }
        for r in rows
    ]


def get_doc_meta(id_doc_rh: int) -> dict | None:
    """Toutes les metadonnees du doc (pour Fen_EditionDocRH)."""
    db = get_pg_connection("rh")
    r = db.query_one(
        """SELECT id_doc_rh, id_type_doc, titre, info_cpl, id_type_produit,
                  doc_actif, prioritaire, id_ste, doc_dpae,
                  doc_dpae_distrib, id_tk_type_photo_dpae,
                  octet_length(contenu) AS taille_contenu
             FROM rh.pgt_doc_rh
            WHERE id_doc_rh = ? LIMIT 1""",
        (int(id_doc_rh),),
    )
    if not r:
        return None
    return {
        "id_doc_rh": str(_int(r.get("id_doc_rh"))),
        "id_type_doc": str(_int(r.get("id_type_doc")) or ""),
        "titre": _str(r.get("titre")),
        "info_cpl": _str(r.get("info_cpl")),
        "id_type_produit": str(_int(r.get("id_type_produit")) or ""),
        "doc_actif": bool(r.get("doc_actif")),
        "prioritaire": bool(r.get("prioritaire")),
        "id_ste": str(_int(r.get("id_ste")) or ""),
        "doc_dpae": bool(r.get("doc_dpae")),
        "doc_dpae_distrib": bool(r.get("doc_dpae_distrib")),
        "id_tk_type_photo_dpae": str(_int(r.get("id_tk_type_photo_dpae")) or ""),
        "taille_contenu": _int(r.get("taille_contenu")),
    }


def create_doc_blank(op_id: int) -> dict:
    """Btn Nouveau : cree un docRH vide. Cf. WinDev HRAZ + HAjoute."""
    db = get_pg_connection("rh")
    new_id = _new_id()
    db.query(
        """INSERT INTO rh.pgt_doc_rh
              (id_doc_rh, id_type_doc, titre, info_cpl, id_type_produit,
               datecrea, doc_actif, prioritaire, id_ste, doc_dpae,
               doc_dpae_distrib, id_tk_type_photo_dpae,
               modif_date, modif_op, modif_elem)
           VALUES (?, NULL, 'Nouveau document', '', 1,
                   NOW(), TRUE, FALSE, 0, FALSE,
                   FALSE, 0,
                   NOW(), ?, 'new')""",
        (new_id, int(op_id)),
    )
    return {"ok": True, "id_doc_rh": str(new_id)}


def update_doc_meta(id_doc_rh: int, payload: dict, op_id: int) -> dict:
    """PUT metadonnees (sans contenu)."""
    db = get_pg_connection("rh")
    db.query(
        """UPDATE rh.pgt_doc_rh
              SET id_type_doc = ?,
                  titre = ?,
                  info_cpl = ?,
                  id_type_produit = ?,
                  id_ste = ?,
                  doc_actif = ?,
                  prioritaire = ?,
                  doc_dpae = ?,
                  doc_dpae_distrib = ?,
                  id_tk_type_photo_dpae = ?,
                  modif_date = NOW(),
                  modif_op = ?,
                  modif_elem = 'modif'
            WHERE id_doc_rh = ?""",
        (
            _int(payload.get("id_type_doc")),
            _str(payload.get("titre")),
            _str(payload.get("info_cpl")),
            _int(payload.get("id_type_produit")) or 1,
            _int(payload.get("id_ste")),
            bool(payload.get("doc_actif")),
            bool(payload.get("prioritaire")),
            bool(payload.get("doc_dpae")),
            bool(payload.get("doc_dpae_distrib")),
            _int(payload.get("id_tk_type_photo_dpae")),
            int(op_id),
            int(id_doc_rh),
        ),
    )
    return {"ok": True}


def upload_doc_content(id_doc_rh: int, content: bytes, op_id: int) -> dict:
    """Replace le bytea 'contenu' (upload d'un .docx)."""
    if not content:
        return {"ok": False, "error": "Contenu vide"}
    db = get_pg_connection("rh")
    import psycopg2
    # On a besoin de Binary pour bytea
    db.query(
        """UPDATE rh.pgt_doc_rh
              SET contenu = ?,
                  modif_date = NOW(),
                  modif_op = ?,
                  modif_elem = 'modif'
            WHERE id_doc_rh = ?""",
        (psycopg2.Binary(content), int(op_id), int(id_doc_rh)),
    )
    return {"ok": True, "taille": len(content)}


def download_doc_content(id_doc_rh: int) -> bytes | None:
    """RETR du bytea 'contenu'.

    Conversion lazy DOCX -> HTML : si le contenu est un DOCX (magic
    PK\\x03\\x04), on le convertit en HTML via mammoth, on persiste le
    HTML en BDD (write-back), puis on retourne le HTML. Cf.
    ctt_ulease.download_doc_content pour la motivation."""
    db = get_pg_connection("rh")
    r = db.query_one(
        "SELECT contenu FROM rh.pgt_doc_rh WHERE id_doc_rh = ? LIMIT 1",
        (int(id_doc_rh),),
    )
    if not r:
        return None
    c = r.get("contenu")
    if c is None:
        return None
    content = bytes(c) if isinstance(c, memoryview) else c
    if content[:4] == b"PK\x03\x04":
        try:
            import io
            import mammoth
            import psycopg2
            html = mammoth.convert_to_html(io.BytesIO(content)).value
            html_bytes = html.encode("utf-8")
            db.query(
                """UPDATE rh.pgt_doc_rh
                      SET contenu = ?
                    WHERE id_doc_rh = ?""",
                (psycopg2.Binary(html_bytes), int(id_doc_rh)),
            )
            return html_bytes
        except Exception:
            return content
    return content


# Donnees fictives pour Publipostage_TESTSalarie (cf. WinDev)
_FAKE_SALARIE = {
    "S_TITRE": "Melle",
    "S_NOM": "NOM_DE_FAMILLE_TRES_LONG",
    "S_PRENOM": "LAURE-EMMANUELLE",
    "S_LNAISS": "Evreux",
    "S_DEPNAISS": "27",
    "S_ADRESSE": "21 rue de la paix",
    "S_CP": "75000",
    "S_VILLE": "PARIS",
    "S_GSM": "0612121212",
    "S_NUMSS": "283092712312355",
    "S_DNAISS": "06/09/1983",
    "FIN_PER_ESSAI": "01/01/2027",
    "DATE_CTS": datetime.now().strftime("%d/%m/%Y"),
    "DATE_ANC": datetime.now().strftime("%d/%m/%Y"),
    "DATE_AVENANT": datetime.now().strftime("%d/%m/%Y"),
    "SECTEURAGENCE": "27.27.27",
    "S_MENTION": "",
    "S_SIGN": "",
}


def _replace_in_docx(content: bytes, variables: dict) -> bytes:
    """Remplace les variables texte dans un DOCX. Utilise python-docx
    qui ne sait pas remplacer si le texte est decoupe sur plusieurs runs ;
    on fusionne les runs par paragraphe a la volee."""
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(content))

    # Tri cles par longueur desc : DATE_AVENANT_FINESSAI avant DATE_AVENANT.
    sorted_vars = sorted(variables.items(), key=lambda kv: -len(kv[0]))

    def replace_in_para(para):
        full = para.text
        modified = full
        for key, val in sorted_vars:
            if key in modified:
                modified = modified.replace(key, str(val))
        if modified != full:
            # Vide tous les runs sauf le premier, met le texte dans le premier
            for run in para.runs[1:]:
                run.text = ""
            if para.runs:
                para.runs[0].text = modified
            else:
                para.add_run(modified)

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)
    # Headers/footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_para(para)
        for para in section.footer.paragraphs:
            replace_in_para(para)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _societe_for_footer(id_ste: int) -> dict:
    """Recupere les infos societe + guimmick base64 pour le footer PDF."""
    db = get_pg_connection("rh")
    s = db.query_one(
        """SELECT raison_sociale, adresse1, cp, ville, siret, guimmick
             FROM rh.pgt_societe WHERE id_ste = ? LIMIT 1""",
        (int(id_ste),),
    )
    if not s:
        return {}
    import base64 as _b64
    g = s.get("guimmick")
    logo_b64 = ""
    if g:
        if isinstance(g, memoryview):
            g = bytes(g)
        if isinstance(g, (bytes, bytearray)):
            sig = bytes(g[:8])
            if sig.startswith(b"\x89PNG"):
                mime = "image/png"
            elif sig.startswith(b"\xff\xd8\xff"):
                mime = "image/jpeg"
            else:
                mime = "image/png"
            logo_b64 = f"data:{mime};base64,{_b64.b64encode(bytes(g)).decode('ascii')}"
    return {
        "raison_sociale": _str(s.get("raison_sociale")),
        "adresse": _str(s.get("adresse1")),
        "cp": _str(s.get("cp")),
        "ville": _str(s.get("ville")),
        "siret": _str(s.get("siret")),
        "logo_b64": logo_b64,
    }


def _societe_variables(id_ste: int) -> dict:
    """Variables STE_* depuis pgt_societe."""
    db = get_pg_connection("rh")
    s = db.query_one(
        """SELECT raison_sociale, code_ape, rcs, capital, adresse1, cp,
                  ville, siren, siret, gerant_nom, gerant_type
             FROM rh.pgt_societe WHERE id_ste = ? LIMIT 1""",
        (int(id_ste),),
    )
    if not s:
        return {}
    return {
        "DOCTITRE": "",
        "STE_RS": _str(s.get("raison_sociale")),
        "STE_APE": _str(s.get("code_ape")),
        "STE_RCS": _str(s.get("rcs")),
        "STE_CAPITAL": _str(s.get("capital")),
        "STE_ADR": f"{_str(s.get('adresse1'))} {_str(s.get('cp'))} {_str(s.get('ville'))}".strip(),
        "STE_VILLE": _str(s.get("ville")),
        "STE_SIREN": _str(s.get("siren")),
        "STE_SIRET": _str(s.get("siret")),
        "STE_GERANT_NOM": _str(s.get("gerant_nom")),
        "STE_GERANT_TYPE": _str(s.get("gerant_type")),
        # Images STE_LOGO / GER_SIGN / STE_CACHET : non gerees en V1.1
        # (necessitent python-docx insertion d'images, V1.2)
        "STE_LOGO": "",
        "GER_SIGN": "",
        "STE_CACHET": "",
    }


def is_docx(content: bytes) -> bool:
    """Detecte le magic header DOCX/ZIP (PK\\x03\\x04)."""
    return bool(content) and content[:4] == b"PK\x03\x04"


def publipostage_test(
    id_doc_rh: int, id_ste: int, titre_doc: str = "",
) -> tuple[bytes, str] | None:
    """Btn 'Tester Mise en page' : substitue les variables S_*/STE_* dans
    le contenu (DOCX ou HTML) et retourne (bytes, mime_type)."""
    content = download_doc_content(id_doc_rh)
    if not content:
        return None
    variables = dict(_FAKE_SALARIE)
    if id_ste:
        ste_vars = _societe_variables(id_ste)
        if titre_doc:
            ste_vars["DOCTITRE"] = titre_doc
        variables.update(ste_vars)

    if is_docx(content):
        try:
            return (
                _replace_in_docx(content, variables),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        except Exception:
            return None

    # HTML : simple str.replace (tri par longueur desc)
    try:
        html = content.decode("utf-8")
        for key, val in sorted(variables.items(), key=lambda kv: -len(kv[0])):
            html = html.replace(key, str(val))
        return html.encode("utf-8"), "text/html; charset=utf-8"
    except Exception:
        return None


def _docx_to_html(docx_bytes: bytes) -> str:
    """Convertit un DOCX en HTML via mammoth (cote serveur)."""
    import io as _io
    import mammoth
    res = mammoth.convert_to_html(_io.BytesIO(docx_bytes))
    return res.value


def _bytes_to_b64_img(img_bytes: bytes | None) -> str:
    """Encode bytes en data: URL (PNG/JPEG/...) pour insertion HTML."""
    if not img_bytes:
        return ""
    import base64 as _b64
    sig = bytes(img_bytes[:8])
    if sig.startswith(b"\x89PNG"):
        mime = "image/png"
    elif sig.startswith(b"\xff\xd8\xff"):
        mime = "image/jpeg"
    elif sig[:6] in (b"GIF87a", b"GIF89a"):
        mime = "image/gif"
    elif sig.startswith(b"BM"):
        mime = "image/bmp"
    else:
        mime = "image/png"
    return f"data:{mime};base64,{_b64.b64encode(bytes(img_bytes)).decode('ascii')}"


def generer_pdf_publiposte(
    body_html: str,
    id_ste: int = 0,
    extra_images: dict[str, str] | None = None,
) -> bytes | None:
    """Genere un PDF avec footer Omaya standard (logo societe / RS+adr+
    SIRET / Page X/Y) via WeasyPrint. Helper reutilisable pour tout
    module qui produit un PDF de doc RH.

    Args:
        body_html : contenu HTML deja publiposte (variables substituees).
        id_ste : pour le footer (guimmick + RS + adresse + SIRET).
        extra_images : substitutions de tokens images supplementaires,
            ex: {'GER_SIGN': 'data:image/png;base64,...',
                 'STE_CACHET': 'data:...'}. STE_LOGO est gere
            automatiquement depuis le guimmick de la societe.

    Renvoie les bytes du PDF ou None si echec.
    """
    import re as _re

    # 1. Footer infos societe (RS / adresse / SIRET / guimmick)
    ste = _societe_for_footer(id_ste) if id_ste else {}
    logo_b64 = ste.get("logo_b64", "")

    # 2. Conversion legacy <font color="..." face="..." size="..."> ->
    # <span style="color:...; font-family:...; font-size:..."> (WeasyPrint).
    def _font_to_span(m: "_re.Match[str]") -> str:
        attrs = m.group(1) or ""
        styles = []
        for prop, attr in (("color", "color"), ("font-family", "face"),
                           ("font-size", "size")):
            val = _re.search(rf'{attr}\s*=\s*"([^"]+)"', attrs, _re.IGNORECASE)
            if val:
                styles.append(f"{prop}: {val.group(1)}")
        if not styles:
            return "<span>"
        return f'<span style="{"; ".join(styles)}">'
    body_html = _re.sub(r"<font\b([^>]*)>", _font_to_span, body_html,
                        flags=_re.IGNORECASE)
    body_html = _re.sub(r"</font>", "</span>", body_html, flags=_re.IGNORECASE)

    # 3. STE_LOGO -> guimmick base64 (si dispo).
    if logo_b64:
        body_html = _re.sub(
            r"STE_LOGO",
            f'<img src="{logo_b64}" '
            f'style="max-height:25mm;max-width:50mm;vertical-align:middle;" '
            f'alt=""/>',
            body_html,
        )
    else:
        body_html = body_html.replace("STE_LOGO", "")

    # 3b. Tokens images supplementaires (GER_SIGN, STE_CACHET, ...).
    if extra_images:
        for token, data_url in extra_images.items():
            if data_url:
                body_html = body_html.replace(
                    token,
                    f'<img src="{data_url}" '
                    f'style="max-height:30mm;max-width:60mm;'
                    f'vertical-align:middle;" alt=""/>',
                )
            else:
                body_html = body_html.replace(token, "")

    # 4. SAUTDEPAGE -> page-break.
    body_html = _re.sub(
        r"(<[^>]+>)?\s*SAUTDEPAGE\s*(</[^>]+>)?",
        '<div style="page-break-before: always;"></div>',
        body_html,
        flags=_re.IGNORECASE,
    )

    # 5. Footer center (RS / adresse / SIRET).
    rs = ste.get("raison_sociale", "")
    adresse = " ".join(
        x for x in [ste.get("adresse", ""),
                    f"{ste.get('cp', '')} {ste.get('ville', '')}".strip()]
        if x
    ).strip()
    siret = ste.get("siret", "")
    footer_center_html = f"<strong>{rs}</strong>" if rs else ""
    if adresse:
        footer_center_html += (
            f"<br/>{adresse}" if footer_center_html else adresse
        )
    if siret:
        footer_center_html += (
            f"<br/>SIRET : {siret}" if footer_center_html else f"SIRET : {siret}"
        )

    # 6. Assemble HTML + CSS @page footer + WeasyPrint.
    html_doc = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
@page {{
    size: A4;
    margin: 25mm 20mm 35mm 20mm;
    @bottom-left {{ content: element(footerLeft); }}
    @bottom-center {{ content: element(footerCenter); }}
    @bottom-right {{
        content: "Page " counter(page) " / " counter(pages);
        font-family: Calibri, "Segoe UI", sans-serif;
        font-size: 9pt;
        color: #000;
    }}
}}
.footerLeft {{ position: running(footerLeft); }}
.footerCenter {{
    position: running(footerCenter);
    font-family: Calibri, "Segoe UI", sans-serif;
    text-align: center;
    font-size: 8pt;
    line-height: 1.3;
    color: #000;
}}
.footerLeft img {{ max-height: 15mm; max-width: 23mm; }}
body {{
    font-family: Calibri, "Segoe UI", sans-serif;
    font-size: 11pt;
    color: #000;
    line-height: 1.4;
}}
table {{ border-collapse: collapse; }}
</style></head><body>
<div class="footerLeft">
{f'<img src="{logo_b64}" alt=""/>' if logo_b64 else ''}
</div>
<div class="footerCenter">
{footer_center_html}
</div>
{body_html}
</body></html>"""

    try:
        from weasyprint import HTML
        return HTML(string=html_doc).write_pdf()
    except Exception:
        return None


def publipostage_test_pdf(
    id_doc_rh: int, id_ste: int, titre_doc: str = "",
) -> bytes | None:
    """Btn 'Tester Mise en page' Fen_EditionDocRH : PDF avec donnees
    fictives + footer auto. Delegue a generer_pdf_publiposte."""
    content = download_doc_content(id_doc_rh)
    if not content:
        return None

    variables = dict(_FAKE_SALARIE)
    if id_ste:
        ste_vars = _societe_variables(id_ste)
        if titre_doc:
            ste_vars["DOCTITRE"] = titre_doc
        variables.update(ste_vars)

    # Convertit DOCX -> HTML si necessaire + substitue les variables.
    if is_docx(content):
        try:
            filled = _replace_in_docx(content, variables)
            body_html = _docx_to_html(filled)
        except Exception:
            return None
    else:
        try:
            body_html = content.decode("utf-8")
            # Tri par longueur desc : DATE_AVENANT_FINESSAI avant DATE_AVENANT.
            for key, val in sorted(
                variables.items(), key=lambda kv: -len(kv[0]),
            ):
                # STE_LOGO est traite par generer_pdf_publiposte.
                if key == "STE_LOGO":
                    continue
                body_html = body_html.replace(key, str(val))
        except Exception:
            return None

    return generer_pdf_publiposte(body_html, id_ste=id_ste)
